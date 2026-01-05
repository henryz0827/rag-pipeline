"""Document loaders for various file types."""
import os
from typing import List, Optional, Dict, Any, Union
from pathlib import Path

from ..vectorstores.base import Document


class DocumentLoader:
    """Universal document loader supporting multiple file types."""

    SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.md', '.docx', '.doc', '.json'}

    def __init__(self, encoding: str = "utf-8"):
        """Initialize document loader.

        Args:
            encoding: Default text encoding.
        """
        self.encoding = encoding

    def load(
        self,
        path: Union[str, Path],
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """Load document(s) from path.

        Args:
            path: File or directory path.
            metadata: Additional metadata to add to documents.

        Returns:
            List of loaded documents.
        """
        path = Path(path)

        if path.is_dir():
            return self.load_directory(path, metadata)
        elif path.is_file():
            return self.load_file(path, metadata)
        else:
            raise FileNotFoundError(f"Path not found: {path}")

    def load_file(
        self,
        file_path: Union[str, Path],
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """Load a single file.

        Args:
            file_path: Path to file.
            metadata: Additional metadata.

        Returns:
            List of documents (usually one).
        """
        file_path = Path(file_path)
        ext = file_path.suffix.lower()

        base_metadata = {
            "source": str(file_path),
            "filename": file_path.name,
            "extension": ext
        }
        if metadata:
            base_metadata.update(metadata)

        if ext == '.pdf':
            return load_pdf(file_path, base_metadata)
        elif ext in {'.txt', '.md'}:
            return load_text(file_path, base_metadata, self.encoding)
        elif ext in {'.docx', '.doc'}:
            return load_docx(file_path, base_metadata)
        elif ext == '.json':
            return load_json(file_path, base_metadata, self.encoding)
        else:
            # Try to load as text
            try:
                return load_text(file_path, base_metadata, self.encoding)
            except Exception:
                raise ValueError(f"Unsupported file type: {ext}")

    def load_directory(
        self,
        dir_path: Union[str, Path],
        metadata: Optional[Dict[str, Any]] = None,
        recursive: bool = True
    ) -> List[Document]:
        """Load all supported documents from a directory.

        Args:
            dir_path: Directory path.
            metadata: Additional metadata.
            recursive: Whether to search subdirectories.

        Returns:
            List of documents.
        """
        dir_path = Path(dir_path)
        documents = []

        pattern = "**/*" if recursive else "*"

        for file_path in dir_path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    docs = self.load_file(file_path, metadata)
                    documents.extend(docs)
                except Exception as e:
                    print(f"Warning: Failed to load {file_path}: {e}")

        return documents


def load_pdf(
    file_path: Union[str, Path],
    metadata: Optional[Dict[str, Any]] = None
) -> List[Document]:
    """Load a PDF file.

    Args:
        file_path: Path to PDF file.
        metadata: Additional metadata.

    Returns:
        List of documents (one per page or combined).
    """
    try:
        import pypdf
    except ImportError:
        try:
            import PyPDF2 as pypdf
        except ImportError:
            raise ImportError(
                "PDF support requires pypdf or PyPDF2. "
                "Install with: pip install pypdf"
            )

    file_path = Path(file_path)
    documents = []

    with open(file_path, 'rb') as f:
        reader = pypdf.PdfReader(f)

        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text.strip():
                page_metadata = {"page": page_num + 1}
                if metadata:
                    page_metadata.update(metadata)

                documents.append(Document(
                    content=text,
                    metadata=page_metadata
                ))

    return documents


def load_text(
    file_path: Union[str, Path],
    metadata: Optional[Dict[str, Any]] = None,
    encoding: str = "utf-8"
) -> List[Document]:
    """Load a text file.

    Args:
        file_path: Path to text file.
        metadata: Additional metadata.
        encoding: File encoding.

    Returns:
        List containing one document.
    """
    file_path = Path(file_path)

    with open(file_path, 'r', encoding=encoding) as f:
        content = f.read()

    return [Document(content=content, metadata=metadata or {})]


def load_docx(
    file_path: Union[str, Path],
    metadata: Optional[Dict[str, Any]] = None
) -> List[Document]:
    """Load a DOCX file.

    Args:
        file_path: Path to DOCX file.
        metadata: Additional metadata.

    Returns:
        List containing one document.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError(
            "DOCX support requires python-docx. "
            "Install with: pip install python-docx"
        )

    file_path = Path(file_path)
    doc = DocxDocument(file_path)

    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    content = '\n\n'.join(paragraphs)

    return [Document(content=content, metadata=metadata or {})]


def load_json(
    file_path: Union[str, Path],
    metadata: Optional[Dict[str, Any]] = None,
    encoding: str = "utf-8"
) -> List[Document]:
    """Load a JSON file with documents.

    Expects JSON to be either:
    - A list of objects with 'content' or 'text' field
    - A single object with 'content' or 'text' field

    Args:
        file_path: Path to JSON file.
        metadata: Additional metadata.
        encoding: File encoding.

    Returns:
        List of documents.
    """
    import json

    file_path = Path(file_path)

    with open(file_path, 'r', encoding=encoding) as f:
        data = json.load(f)

    documents = []

    if isinstance(data, list):
        for item in data:
            content = item.get('content') or item.get('text') or str(item)
            item_metadata = {k: v for k, v in item.items() if k not in {'content', 'text'}}
            if metadata:
                item_metadata.update(metadata)

            documents.append(Document(content=content, metadata=item_metadata))
    elif isinstance(data, dict):
        content = data.get('content') or data.get('text') or str(data)
        item_metadata = {k: v for k, v in data.items() if k not in {'content', 'text'}}
        if metadata:
            item_metadata.update(metadata)

        documents.append(Document(content=content, metadata=item_metadata))

    return documents
